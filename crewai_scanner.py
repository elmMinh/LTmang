import nmap3
import requests
import pandas as pd
import concurrent.futures
import smtplib
from email.message import EmailMessage
from docx import Document
import datetime
import os

# Hàm gọi DeepSeek AI để phân tích
def ai_analyze(text):
    url = "https://api.deepseek.com/v1/chat/completions"
    headers = {
        "Authorization": "Bearer your_deepseek_api_key",  # Thay bằng API KEY của bạn
        "Content-Type": "application/json"
    }
    payload = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": "Bạn là chuyên gia an ninh mạng, hãy phân tích kết quả quét lỗ hổng."},
            {"role": "user", "content": text}
        ],
        "temperature": 0.3
    }
    try:
        response = requests.post(url, headers=headers, json=payload)
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"DeepSeek AI Error: {e}")
        return "AI Analysis Failed"

# Network Scanner Agent
class NetworkScannerAgent:
    def __init__(self):
        self.scanner = nmap3.Nmap()
    
    def scan(self, target_ip):
        print(f"[NetworkScanner] Scanning {target_ip}...")
        result = self.scanner.scan_top_ports(target_ip)
        return result

# Web Scanner Agent
class WebScannerAgent:
    def scan_http(self, target_ip):
        urls = [f"http://{target_ip}", f"https://{target_ip}"]
        results = []
        for url in urls:
            try:
                r = requests.get(url, timeout=5)
                results.append((url, r.status_code))
            except Exception as e:
                results.append((url, str(e)))
        return results

# Analysis Agent
class AnalysisAgent:
    def __init__(self):
        self.reports = []
    
    def analyze(self, network_results, web_results):
        report = ""
        for ip, res in network_results.items():
            report += f"\nHost: {ip}\nOpen Ports:\n"
            ports = res.get("ports", [])
            for p in ports:
                portid = p['portid']
                name = p['service']['name']
                report += f"- Port {portid}: {name}\n"
        for web in web_results:
            report += f"\nWeb Check: {web[0]} - Status: {web[1]}\n"
        
        ai_summary = ai_analyze(report)
        final_report = report + "\n\nAI Analysis:\n" + ai_summary
        self.reports.append(final_report)
        return final_report

    def export_report_txt(self, filename):
        with open(filename, "w", encoding="utf-8") as f:
            for rep in self.reports:
                f.write(rep)
    
    def export_summary_excel(self, filename):
        data = []
        for rep in self.reports:
            sections = rep.split("Host:")
            for sec in sections[1:]:
                lines = sec.strip().splitlines()
                ip = lines[0]
                ports = []
                for l in lines[1:]:
                    if l.startswith("- Port"):
                        parts = l.split(":")
                        port = parts[0].replace("- Port", "").strip()
                        service = parts[1].strip()
                        ports.append(f"{port}/{service}")
                data.append({"Host": ip, "Open Ports": ", ".join(ports)})
        df = pd.DataFrame(data)
        df.to_excel(filename, index=False)

    def export_report_docx(self, filename):
        doc = Document()
        doc.add_heading('CrewAI Vulnerability Scanner Report', 0)
        doc.add_paragraph(f"Scan Time: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        for rep in self.reports:
            sections = rep.split("Host:")
            for sec in sections[1:]:
                lines = sec.strip().splitlines()
                ip = lines[0]
                ports = []
                for l in lines[1:]:
                    if l.startswith("- Port"):
                        parts = l.split(":")
                        port = parts[0].replace("- Port", "").strip()
                        service = parts[1].strip()
                        ports.append(f"{port}/{service}")
                table = doc.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'IP'
                hdr_cells[1].text = 'Open Ports'
                row_cells = table.add_row().cells
                row_cells[0].text = ip
                row_cells[1].text = ", ".join(ports)
                doc.add_paragraph()
        doc.save(filename)

# Manager Agent
class ManagerAgent:
    def __init__(self, targets):
        self.targets = targets
        self.network_agent = NetworkScannerAgent()
        self.web_agent = WebScannerAgent()
        self.analysis_agent = AnalysisAgent()

    def scan_target(self, target):
        net_res = self.network_agent.scan(target)
        web_res = self.web_agent.scan_http(target)
        return target, net_res.get(target, {}), web_res

    def send_email_report(self, txt_file, excel_file, docx_file):
        smtp_server = input("SMTP Server: ")
        smtp_port = int(input("SMTP Port (587 or 465): "))
        email_sender = input("Your Email: ")
        email_pass = input("Email Password: ")
        email_receiver = input("Receiver Email: ")

        msg = EmailMessage()
        msg['Subject'] = 'CrewAI Scanner Report'
        msg['From'] = email_sender
        msg['To'] = email_receiver
        msg.set_content('Attached: Scan Reports')

        for file in [txt_file, excel_file, docx_file]:
            with open(file, 'rb') as f:
                msg.add_attachment(f.read(), maintype='application', subtype='octet-stream', filename=os.path.basename(file))
        
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(email_sender, email_pass)
            server.send_message(msg)
        print("Email Sent Successfully!")

    def execute(self):
        network_results = {}
        web_results = []
        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
            futures = [executor.submit(self.scan_target, t) for t in self.targets]
            for future in concurrent.futures.as_completed(futures):
                t, net, web = future.result()
                network_results[t] = net
                web_results.extend(web)
        
        report = self.analysis_agent.analyze(network_results, web_results)

        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        txt_file = f"reports/scan_report_{timestamp}.txt"
        excel_file = f"reports/scan_summary_{timestamp}.xlsx"
        docx_file = f"reports/scan_report_{timestamp}.docx"

        os.makedirs("reports", exist_ok=True)

        self.analysis_agent.export_report_txt(txt_file)
        self.analysis_agent.export_summary_excel(excel_file)
        self.analysis_agent.export_report_docx(docx_file)

        self.send_email_report(txt_file, excel_file, docx_file)
        print("\nScan Completed and Reports Generated.")

if __name__ == "__main__":
    targets = []
    while True:
        t = input("Enter target IP/domain (leave blank to start scanning): ")
        if not t:
            break
        targets.append(t)
    if targets:
        manager = ManagerAgent(targets)
        manager.execute()
    else:
        print("No targets entered. Exiting.")
